#!/usr/bin/python3

from PyQt5.QtWidgets import *
from PyQt5.QtGui import *
from PyQt5.QtCore import QTimer
import ctypes
import winsound
from stat import S_IREAD, S_IRGRP, S_IROTH, S_IWUSR
from ctypes import wintypes
import traceback
from bs4 import BeautifulSoup as bs
from urllib.parse import urlparse, urlencode
import smtplib
from email.mime.text import MIMEText
import pandas as pd
import fileinput
from ldap3 import Server, Connection, ALL, ALL_ATTRIBUTES
from pandas import *
# import docx
# from docx import exceptions
# import xlrd
from docx.shared import Pt
from docx.enum.section import *
from docx.enum.text import *
from docxtpl import Document
import os
import pyodbc
import json
import socket
import datetime
import requests
import re

val = False
run = True
u=''
user=''
loc=''
num=''
scan = ''
email = ''
log = "H:\Informatics\Pharmacy Equipment\\redirx-test.log"
sdt = datetime.datetime.now()
w = socket.gethostname()
ws = w.replace('.mskcc.org', '')

class loginWindow(QDialog):
    def __init__(self,parent=None):
        super(loginWindow,self).__init__(parent)
        self.setWindowTitle('ReDi-Rx')
        self.loginUser = QLineEdit()
        self.loginPass = QLineEdit()
        self.loginPass.setEchoMode(QLineEdit.Password)
        self.loginULabel = QLabel("Outlook ID:")
        self.loginPLabel = QLabel("Password:")
        self.loginButton = QPushButton('Login',self)
        self.loginButton.clicked.connect(self.handleLogin)
        llayout = QGridLayout(self)
        llayout.addWidget(self.loginULabel,1,1,1,1)
        llayout.addWidget(self.loginUser,1,2,1,1)
        llayout.addWidget(self.loginPLabel,2,1,1,1)
        llayout.addWidget(self.loginPass,2,2,1,1)
        llayout.addWidget(self.loginButton,3,4,1,1)

    def handleLogin(self):
        global u
        global user
        global val
        global admin
        global loc
        global num
        global email
        u = self.loginUser.text()
        p = self.loginPass.text()
        locnum = pd.read_excel("H:\Informatics\Pharmacy Equipment\WS.xlsx", sheet_name="Pharm WS Info")
        if locnum.loc[locnum['WSID'] == ws, 'WSID'].empty:
            # QMessageBox.warning(self, 'Warning',
            #                     'This is not a registered Pharmacy workstation.\nPlease submit the workstation information to the Pharmacy Informatics team.',
            #                     QMessageBox.Ok, QMessageBox.Ok)
            loc = 'Unassigned WS'
            num = '999-999-9999'
            email = '@mskcc.org'
        else:
            if locnum.loc[locnum['WSID'] == ws, 'email'].empty:
                email = '@mskcc.org'
            else:
                email = locnum.loc[locnum['WSID'] == ws, 'email'].iloc[0]
            if locnum.loc[locnum['WSID'] == ws, 'Contact#'].empty:
                num = '999-999-9999'
            else:
                num = locnum.loc[locnum['WSID'] == ws, 'Contact#'].iloc[0]
            if locnum.loc[locnum['WSID'] == ws, 'Area'].empty:
                loc = 'Unassigned WS'
            else:
                loc = locnum.loc[locnum['WSID'] == ws, 'Area'].iloc[0]

        if p != '':
            server = Server(host='MSKCC.ROOT.MSKCC.ORG', use_ssl=True, get_info=ALL)
            conn = Connection(server, user='MSKCC\\' + u, password=p)
            p = ''
            if not conn.bind():
                QMessageBox.warning(self, "Warning", "Invalid Credentials.", QMessageBox.Ok, QMessageBox.Ok)
                conn.password = ''
                val = False

            else:
                self.accept()
                val = True
                Base = 'dc=mskcc,dc=root,dc=mskcc,dc=org'
                conn.search(search_base=Base, search_filter='(cn=' + u + ')', attributes=[ALL_ATTRIBUTES])
                cd = json.loads(conn.response_to_json())
                if cd['entries'][0]['attributes'] == 'extensionAttribute15':
                    user = cd['entries'][0]['attributes']['extensionAttribute15']
                else:
                    user = cd['entries'][0]['attributes']['givenName']+ " " + cd['entries'][0]['attributes']['sn']
                try:
                    d = cd['entries'][0]['attributes']['memberOf']
                    if 'CN=GRP_PHA_Informatics,OU=ezGroups,OU=Resources,DC=MSKCC,DC=ROOT,DC=MSKCC,DC=ORG' in d:
                        admin = 'Y'
                    else:
                        admin = 'N'
                except Exception as e:
                    admin = 'N'

                conn.unbind()
                conn.password = ''
                p = ''
                os.chmod(log, S_IWUSR | S_IREAD)
                logw = open(log, "a+")
                logw.write("\nSession started at " + str(sdt) + " by user " + user + " from Workstation " + ws + ", located at " + loc + "; Phone: " + str(num) + ";email: " + email + "\n")
                logw.close()
                os.chmod(log, S_IREAD | S_IRGRP | S_IROTH)

        else:
            QMessageBox.warning(self, "Warning", "Your Password must be entered.", QMessageBox.Ok)
            val = False

class mainWindow(QMainWindow):
    def __init__(self):
        super().__init__()
        self.title = 'ReDi-Rx'
        self.left = 200
        self.top = 200
        self.width = 1000
        self.height = 200
        self.setWindowTitle(self.title)
        self.setGeometry(self.left,self.top,self.width,self.height)
        self.tableWidget = initUi(self)
        self.setCentralWidget(self.tableWidget)
        self.setWindowIcon(QIcon('icon.png'))

class initUi(QWidget): #setting up UI elements#
    def __init__(self, parent=None):
        super(QWidget, self).__init__(parent)
        self.serialEnabled = False
        self.layout = QGridLayout(self)
        self.user = QLabel()
        self.loc = QLineEdit()
        self.phone = QLineEdit()
        global log
        global dt
        global ws
        global val
        global run
        global u
        global admin
        global loc
        global num

        self.loc.setText(loc)
        self.phone.setText(str(num))

        self.user.setText(u + " / " + ws)
        self.tabs = QTabWidget()
        self.tab1 = QWidget()
        self.tab2 = QWidget()
        self.tab3 = QWidget()
        self.tab4 = QWidget()
        self.tab5 = QWidget()
        self.tab6 = QWidget()
        self.tab7 = QWidget()
        self.tabs.addTab(self.tab1, "Receiving")
        self.tabs.addTab(self.tab2, "Dispensing")  # This will be phased-in in 2020 #
        self.tabs.addTab(self.tab6, "PS80 Check")
        self.tabs.addTab(self.tab7, "Rejection Reporting")
        if admin == 'Y':
            self.tabs.addTab(self.tab3, "KBMA Requests")
            self.tabs.addTab(self.tab4, "Reporting")
            self.tabs.addTab(self.tab5, "Lot/Exp Maintenance")
        self.tab1.layout = QHBoxLayout(self)
        self.tab2.layout = QGridLayout(self)
        self.tab3.layout = QHBoxLayout(self)
        self.tab4.layout = QHBoxLayout(self)
        self.tab5.layout = QHBoxLayout(self)
        self.tab6.layout = QHBoxLayout(self)
        self.tab7.layout = QGridLayout(self)

        # Setting up Receiving Tab #
        search_label = QLabel("Scan:")
        self.scan = QLineEdit()
        self.doseN = QLineEdit()
        self.doseQ = QLineEdit()
        self.packQ = QLineEdit()
        self.scan.returnPressed.connect(self.parseScan)
        srx_result = QLabel("SRx Match:")
        self.image_res = QLabel()
        drug = QLabel("Drug:")
        stg = QLabel("Strength:")
        brand = QLabel("Brand:")
        manf = QLabel("MFR:")
        ndc = QLabel("NDC:")
        desc = QLabel("Package:")
        rte = QLabel("Route:")
        form = QLabel("Form:")
        lotl = QLabel("Lot:")
        expl = QLabel("Exp:")

        self.srx = QLabel()
        self.str = QLineEdit()
        self.str.clear()
        self.route = QLineEdit()
        self.route.clear()
        self.bname = QLineEdit()
        self.bname.clear()
        self.mfr = QLineEdit()
        self.mfr.clear()
        self.ndc = QLineEdit()
        self.ndc.clear()
        self.drug = QLineEdit()
        self.drug.clear()
        self.pack = QLineEdit()
        self.pack.clear()
        self.dform = QLineEdit()
        self.dform.clear()
        self.lot = QLineEdit()
        self.lot.clear()
        self.exp = QLineEdit()
        self.exp.clear()
        self.lexp = QLineEdit()
        self.lexp.clear()
        self.sndc = QLineEdit()
        self.sndc.clear()
        self.scans1 = QLineEdit()
        self.scans1.clear()
        self.srxd = QLineEdit()
        self.srxm = QLineEdit

        self.search1 = QPushButton("Get Info")
        self.search1.setAutoDefault(True)
        self.search1.clicked.connect(self.searchButton)  # this connects the clicking action to a function(method) below
        self.search1.setEnabled(False)

        self.submitR = QPushButton("Submit")
        self.submitR.clicked.connect(self.submitReq)
        self.submitR.setEnabled(False)

        self.exit = QPushButton("Exit")
        self.exit.clicked.connect(self.closeApp)

        self.tab1.layout.addWidget(search_label)
        self.tab1.layout.addWidget(self.scan)
        self.tab1.layout.addWidget(self.image_res)
        self.tab1.layout.addWidget(self.search1)
        self.tab1.layout.addWidget(self.submitR)

        # setting up the grid view below #
        self.h1 = QGridLayout(self)
        self.h1.addLayout(self.tab1.layout, 1, 0, 1, 6)
        self.h1.addWidget(srx_result, 3, 0, 1, 1)
        self.h1.addWidget(self.srx, 3, 1, 1, 6)
        self.h1.addWidget(drug, 14, 0, 1, 1)
        self.h1.addWidget(self.drug, 14, 1, 1, 1)
        self.h1.addWidget(stg, 14, 2, 1, 1)
        self.h1.addWidget(self.str, 14, 3, 1, 1)
        self.h1.addWidget(rte, 16, 0, 1, 1)
        self.h1.addWidget(self.route, 16, 1, 1, 1)
        self.h1.addWidget(brand, 16, 2, 1, 1)
        self.h1.addWidget(self.bname, 16, 3, 1, 1)
        self.h1.addWidget(ndc, 17, 2, 1, 1)
        self.h1.addWidget(self.ndc, 17, 3, 1, 1)
        self.h1.addWidget(manf, 17, 0, 1, 1)
        self.h1.addWidget(self.mfr, 17, 1, 1, 1)
        self.h1.addWidget(desc, 18, 0, 1, 1)
        self.h1.addWidget(self.pack, 18, 1, 1, 1)
        self.h1.addWidget(form, 18, 2, 1, 1)
        self.h1.addWidget(self.dform, 18, 3, 1, 1)
        self.h1.addWidget(lotl, 19, 0, 1, 1)
        self.h1.addWidget(self.lot, 19, 1, 1, 1)
        self.h1.addWidget(expl, 19, 2, 1, 1)
        self.h1.addWidget(self.exp, 19, 3, 1, 1)
        self.tab1.setLayout(self.h1)

        # Setting up Dispensing Tab #
        self.scan2 = QLineEdit()
        self.lexpi = QLineEdit()
        self.lastScan = QLineEdit()
        self.scan3 = QLineEdit()
        self.res = QLineEdit()
        self.lexp2 = QLineEdit()
        self.scan2.returnPressed.connect(self.parseScan2)
        disp = QLabel("Verify Scan:")
        match2 = QLabel('CIS Order:')
        product2 = QLabel('Product Scanned:')
        self.srx2 = QLabel()
        self.srx2.clear()
        self.srp2 = QLabel()
        self.srp2.clear()
        ndcl = QLabel('NDC:')
        self.ndc2 = QLineEdit()
        self.ndc2.clear()
        lotl2 = QLabel('Lot Number:')
        self.lot2 = QLineEdit()
        self.lot2.clear()
        self.submitD = QPushButton("Submit")
        self.submitD.clicked.connect(self.submitDisp)
        self.ndcCount = 0
        self.ndcLine = 4
        self.ndcs = []

        expl2 = QLabel('Exp:')
        self.addItem = QPushButton("+")
        self.addItem.clicked.connect(self.addMoreItems)
        self.exp2 = QLineEdit()
        self.exp2.clear()
        self.image_res2 = QLabel()
        self.f = QGridLayout()
        self.f.addLayout(self.tab2.layout, 0, 0, 1, 1)
        self.f.addWidget(disp, 1, 0, 1, 1)
        self.f.addWidget(self.scan2, 1, 1, 1, 1)
        self.f.addWidget(self.image_res2, 1, 2, 1, 3)
        self.f.addWidget(match2, 2, 0, 1, 1)
        self.f.addWidget(self.srx2, 2, 1, 1, 5)
        self.f.addWidget(product2, 3, 0, 1, 1)
        self.f.addWidget(self.srp2, 3, 1, 1, 3)
        self.f.addWidget(ndcl, 4, 0, 1, 1)
        self.f.addWidget(self.ndc2, 4, 1, 1, 1)
        self.f.addWidget(lotl2, 4, 2, 1, 1)
        self.f.addWidget(self.lot2, 4, 3, 1, 1)
        self.f.addWidget(expl2, 4, 4, 1, 1)
        self.f.addWidget(self.exp2, 4, 5, 1, 1)
        self.f.addWidget(self.addItem, 4, 6, 1, 1)
        self.f.addWidget(self.submitD, 10, 6, 1, 1)
        self.tab2.setLayout(self.f)

        # Setting up Reporting Tab #
        self.report = QPushButton("Print Report")
        # report = Reporting.Dispensing(self)
        self.report.clicked.connect(self.Dispensed)
        r = QGridLayout()
        r.addLayout(self.tab4.layout, 1, 1, 1, 1)
        r.addWidget(self.report, 2, 1, 1, 1)

        self.tab4.setLayout(r)

        # Setting Up PI Tab #
        self.piTable = QTableWidget()
        self.piTable.setColumnCount(14)
        self.getPI = QPushButton('See Requests')
        self.getPI.clicked.connect(self.getPi)
        self.tab3.layout.addWidget(self.getPI)
        self.submitPi = QPushButton('Complete Request')
        # self.submitPi.setEnabled(False)
        self.submitPi.clicked.connect(self.submitPI)
        self.tab3.layout.addWidget(self.submitPi)
        h2 = QGridLayout(self)
        h2.addLayout(self.tab3.layout, 1, 1, 1, 1)
        h2.addWidget(self.piTable, 2, 1, 1, 6)
        self.tab3.setLayout(h2)

        # Setting up Lot/Exp Maintenance Tab #
        newlexp = QLabel("New Item:")
        self.newLexp = QLineEdit()
        self.lexpt = QTableWidget()
        self.lexpt.setColumnCount(1)
        # This sets up the list of current items#
        lines = []
        with open('lexp.txt', 'r') as reader:
            line = reader.readlines()

            for item in line:
                item.replace('\n', '')
                # print(item.rstrip())
                lines.append(item.rstrip())
            # print(lines)
            self.lexpt.setHorizontalHeaderLabels(["SCM Item"])
            self.lexpt.setRowCount(len(lines))

            for line in enumerate(lines):
                # data = line[1]
                for column, data in enumerate(line):
                    # print(line)
                    self.lexpt.setItem(line[0], column, QTableWidgetItem(str(line[1])))
        # This prevents edit from the table itself#

        self.lexpt.resizeColumnsToContents()
        self.lexpt.setEditTriggers(QAbstractItemView.NoEditTriggers)
        self.addLexp = QPushButton('Add to List')
        self.remLexp = QPushButton('Remove from List')
        self.addLexp.clicked.connect(self.add2Lexp)
        self.remLexp.clicked.connect(self.rem2Lexp)
        self.tab5.layout.addWidget(newlexp)
        self.tab5.layout.addWidget(self.newLexp)
        self.tab5.layout.addWidget(self.addLexp)
        self.tab5.layout.addWidget(self.remLexp)
        h5 = QGridLayout(self)
        h5.addLayout(self.tab5.layout, 1, 1, 1, 1)
        h5.addWidget(self.lexpt, 2, 1, 1, 2)
        self.tab5.setLayout(h5)

        # PS 80 CHECK TAB #
        self.psLabel = QLabel("Scan Product:")
        self.psScan = QLineEdit()
        self.psResult = QLabel()
        self.psImage = QLabel()
        self.psDiscIm = QLabel()
        di = QPixmap("disc.jpg").scaledToHeight(25)
        self.psDiscIm.setPixmap(di)
        self.psDisc = QLabel("DISCLAIMER: The information on polysorbate 80 is from the DailyMed database. It is provided solely for information purposes.\nUsers of this information should always use their own clinical discretion.")
        self.psDisc.setFont(QFont('Verdana', 12))
        self.psScan.returnPressed.connect(self.parseScan3)
        self.tab6.layout.addWidget(self.psLabel)
        self.tab6.layout.addWidget(self.psScan)
        self.tab6.layout.addWidget(self.psImage)
        h6 = QGridLayout(self)
        h6.addLayout(self.tab6.layout, 2, 2, 1, 1)
        h6.addWidget(self.psDiscIm, 1, 1, 1, 1)
        h6.addWidget(self.psDisc, 1, 2, 1, 1)
        h6.addWidget(self.psResult, 3, 2, 1, 1)
        self.tab6.setLayout(h6)

        # Rejection Reporting Tab #

        self.rejLabel = QLabel("Scan Label:")
        self.rejScan = QLineEdit()
        self.rejScan.returnPressed.connect(self.rejFind)
        self.rejResult = QLabel()
        self.rejScan.setFixedSize(150, 22)
        self.rejResult.setWordWrap(True)
        self.reasonLabel = QLabel("Reason:")
        self.reasonDrop = QComboBox()
        self.reasonDrop.addItems(["Expired Med", "Wrong Diluent Type", "Wrong Dose/Volume", "Wrong Drug", "Wrong Patient", "Wrong Reconstruction", "Other"])
        self.rejComLabel = QLabel("Comments:")
        self.rejComField = QTextEdit()
        self.rejSub = QPushButton("Submit Rejection")
        self.rejSub.clicked.connect(self.rejSubmit)
        self.rejSub.setEnabled(False)
        self.rejComField.setFixedSize(300, 88)
        self.rejLotLabel = QLabel("Batch Lot:")
        self.rejLotField = QLineEdit()
        self.rejLotField.setFixedSize(150, 22)
        self.rejExpLabel = QLabel("Batch Exp:")
        self.rejExpField = QLineEdit()
        self.rejExpField.setFixedSize(150, 22)
        self.rejDesc = QLineEdit()
        self.rejDoseQty = QLineEdit()
        self.rejDoseUom = QLineEdit()
        self.rejPtLoc = QLineEdit()
        self.rejType = QLabel()
        self.scan4 = QLineEdit()
        self.rdoseNum = QLineEdit()
        self.tab7.layout.addWidget(self.rejLabel, 0, 0, 1, 1)
        self.tab7.layout.addWidget(self.rejScan, 0, 1, 1, 1)
        self.tab7.layout.addWidget(self.rejResult, 0, 2, 1, 14)
        self.h7 = QGridLayout(self)
        self.h7.addLayout(self.tab7.layout, 1, 0, 1, 16)
        self.h7.addWidget(self.rejLotLabel, 10, 0, 1, 1)
        self.h7.addWidget(self.rejLotField, 10, 1, 1, 1)
        self.h7.addWidget(self.rejExpLabel, 10, 2, 1, 1)
        self.h7.addWidget(self.rejExpField, 10, 3, 1, 1)

        self.h7.addWidget(self.reasonLabel, 12, 0, 1, 1)
        self.h7.addWidget(self.reasonDrop, 12, 1, 1, 1)
        self.h7.addWidget(self.rejComLabel, 14, 0, 1, 1)
        self.h7.addWidget(self.rejComField, 14, 1, 1, 15)
        self.h7.addWidget(self.rejSub, 16, 0, 1, 2)
        self.rejLotLabel.setVisible(False)
        self.rejLotField.setVisible(False)
        self.rejExpLabel.setVisible(False)
        self.rejExpField.setVisible(False)
        self.tab7.setLayout(self.h7)
        self.rejScan.setFocus()

        self.layout.addWidget(self.tabs, 0, 0, 1, 6)
        self.layout.addWidget(self.user, 6, 0, 1, 1)
        self.layout.addWidget(self.exit, 6, 5, 1, 1)
        self.scan.setFocus()  # bring cursor to scan field upon start#
        self.show()

        self.timer = QTimer()
        self.timer.start(14400000)
        self.timer.timeout.connect(self.closeApp)

    def closeApp(self):
        global run
        global val
        val = False
        #app.exec()
        sys.exit()
        #app.exec_()

    def addMoreItems(self):
        if len(self.ndc2.text()) == 0:
            QMessageBox.warning(self, "Warning", "You must scan at least one product before adding more.",QMessageBox.Ok, QMessageBox.Ok)
            self.scan2.selectAll()
        elif self.lastScan.text() == "F":
            QMessageBox.warning(self, "Warning", "You have an Incorrect Scan.\nPlease review before adding more items.",QMessageBox.Ok, QMessageBox.Ok)
        # below lines for capturing for Marie Ryan's list (Albumin, Blood Prods, IVIGs)#
        elif self.f.count() == 16 and (len(self.lot2.text()) == 0 or len(self.exp2.text()) == 0):
            QMessageBox.warning(self, "Warning", "You must record all information before adding more products.",QMessageBox.Ok, QMessageBox.Ok)
            self.scan2.selectAll()

        elif self.f.count() == 22 and len(self.f.itemAt(17).widget().text()) == 0:
            QMessageBox.warning(self, "Warning", "You must scan at least one product before adding more.",QMessageBox.Ok, QMessageBox.Ok)
            self.scan2.selectAll()
        elif self.f.count() == 22 and (len(self.f.itemAt(17).widget().text()) != 0 and (len(self.f.itemAt(19).widget().text()) == 0 or len(self.f.itemAt(21).widget().text()) == 0)):
            QMessageBox.warning(self, "Warning", "You must record all information before adding more products.",QMessageBox.Ok, QMessageBox.Ok)
            self.scan2.selectAll()
        elif self.ndcLine == 6:
            QMessageBox.warning(self, "Warning", "No more than 3 different Lots can be recorded.",QMessageBox.Ok,QMessageBox.Ok)
            self.scan2.selectAll()
        else:
            i = self.ndcCount
            line = self.ndcLine
            line += 1
            self.ndcs.append(QLabel("NDC:"))
            self.ndcs.append(QLineEdit(self))

            self.ndcs.append(QLabel("Lot Number:"))
            self.ndcs.append(QLineEdit(self))

            self.ndcs.append(QLabel("Exp:"))
            self.ndcs.append(QLineEdit(self))

            self.ndcs[line].setLayout = QGridLayout()
            self.f.addWidget(self.ndcs[i], line, 0, 1, 1)
            self.f.addWidget(self.ndcs[i + 1], line, 1, 1, 1)
            self.f.addWidget(self.ndcs[i + 2], line, 2, 1, 1)
            self.f.addWidget(self.ndcs[i + 3], line, 3, 1, 1)
            self.f.addWidget(self.ndcs[i + 4], line, 4, 1, 1)
            self.f.addWidget(self.ndcs[i + 5], line, 5, 1, 1)
            #print(self.ndcCount, self.f.indexOf(self.ndc2), self.f.indexOf(self.ndcs[i+1]), self.f.count())
            self.ndcCount += 6
            self.ndcLine += 1
        self.scan2.setFocus()
        self.scan2.selectAll()
        self.timer.start(14400000)

    def searchButton(self):
        global run
        global log
        global ndc
        pck = ''
        mfr = ''
        rt= ''
        st = ''
        prop = ''
        self.srx.setText('')
        self.image_res.clear()
        sv = self.scans1.text()
        nm = QPixmap("NoMatch.png").scaledToHeight(25)
        m = QPixmap("Match.png").scaledToHeight(25)
        if sv == '':
            QMessageBox.warning(self, 'Warning', "Please Scan a Product First.", QMessageBox.Ok, QMessageBox.Ok)
        else:
            os.chmod(log, S_IWUSR | S_IREAD)
            logw = open(log, 'a+')
            try:  # API Calls to DailyMed, OpenFDA, NIH #
                params = {"labeltype": "all", "query": sv}
                r = requests.get("https://dailymed.nlm.nih.gov/dailymed/search.cfm", params=params)
                o = urlparse(r.url).query
                # print(o)
                soup = bs(r.text, "html5lib")
                if o.startswith('setid='):
                    q = o.replace("setid=", '')
                else:
                    mdt = str(soup.find("a", {"class": "drug-info-link"}))
                    mdt = mdt[mdt.find("setid=") + 6:mdt.find("\">")]
                    # print(mdt)
                    if mdt == '':
                        logw.write(
                            'Error Code 10: No Drug Package Information found in OpenFDA.\n')
                        QMessageBox.warning(self, 'Warning', "No Information found for this Scancode.\nPlease Enter Information Manually.",
                                            QMessageBox.Ok, QMessageBox.Ok)
                        q = ""
                        self.image_res.setPixmap(nm)
                        self.image_res.setMaximumWidth(50)
                    else:
                        q = mdt.replace("setid=", '')

                if q != "":
                    params = {'spl_id': '"%s"' % q}
                    p = urlencode(params)
                    r = requests.get('https://api.fda.gov/drug/ndc.json?search=%s' % p).text
                    # print(r)
                    res = json.loads(r)
                    results = res['results'][0]
                    # print('From FDA:\n',json.dumps(results, indent=3))
                    self.drug.setText(results['generic_name'])
                    self.bname.setText(results['brand_name'])
                    param = {'id': q}
                    r2 = requests.get('https://rxnav.nlm.nih.gov/REST/ndcproperties.json', params=param).text
                    d2 = json.loads(r2)
                    # print('From NIH ndc properties:\n',json.dumps(d2, indent=3))
                    if not d2:
                        self.ndc.setText(results['packaging'][0]['package_ndc'])
                        self.pack.setText(results['packaging'][0]['description'])
                        self.mfr.setText(results['labeler_name'])
                        self.dform.setText(results['dosage_form'])
                        l = 0
                        for l in range(len(results['route'])):
                            rt += results['route'][l] + ', '
                        # print(rt)
                        self.route.setText(rt.rstrip(', '))
                        l = 0
                        for l in range(len(results['active_ingredients'])):
                            st += results['active_ingredients'][l]['strength'] + ', '
                        self.str.setText(st.rstrip(', '))
                    else:
                        prop = d2['ndcPropertyList']['ndcProperty']
                        for i in range(len(prop)):
                            if self.scans1.text() == str(prop[i]['ndc10']).replace('-', ''):
                                # print(prop[i]['rxcui'])
                                rxcui = prop[i]['rxcui']
                                self.ndc.setText(prop[i]['ndc10'])
                                j = 0
                                for j in range(len(prop[i]['packagingList']['packaging'])):
                                    pck += prop[i]['packagingList']['packaging'][j] + ', '
                                self.pack.setText(pck.rstrip(', '))
                                k = 0
                                pp = prop[i]['propertyConceptList']['propertyConcept']
                                for k in range(len(prop[i]['propertyConceptList']['propertyConcept'])):
                                    if pp[k]['propName'] == "LABELER":
                                        mfr += pp[k]['propValue'] + ', '
                                self.mfr.setText(mfr.rstrip(', '))
                                # print(pp[k]['propValue'])

                        r3 = requests.get('https://rxnav.nlm.nih.gov/REST/RxTerms/rxcui/%s/allinfo.json' % rxcui).text
                        d3 = json.loads(r3)
                        # print('From nih rxcui info:\n',json.dumps(d3, indent=3))

                        self.str.setText(d3['rxtermsProperties']['strength'])
                        self.dform.setText(d3['rxtermsProperties']['rxnormDoseForm'])
                        self.route.setText(d3['rxtermsProperties']['route'])
                        q = ''

                    self.srx.setText('NOT IN SRX')
                    self.image_res.setPixmap(nm)
                    self.image_res.setMaximumWidth(50)
                    logw.write("API Call information for Scan {}: {} by {}; package: {}, NDC {}\n".format(
                        scan, results['generic_name'], self.mfr.text(),
                        self.pack.text(), prop[i]['ndc10']))
                else:
                    pass
            except Exception as e:
                logw.write('Error GetInfo: %s' % e)
            self.submitR.setEnabled(True)
            logw.close()
        self.timer.start(14400000)

    def submitReq(self):
        global u
        global user
        global email
        global log
        loc = self.loc.text()
        num = self.phone.text().strip()
        if len(self.sndc.text())==0 or len(self.drug.text())==0 or len(self.mfr.text())==0\
            or len(self.bname.text())==0 or len(self.str.text())==0\
            or len(self.dform.text())==0 or len(self.pack.text())==0 or len(self.scan.text())==0:
            QMessageBox.warning(self, "Warning", "All fields must have a value.", QMessageBox.Ok, QMessageBox.Ok)
            self.scan2.selectAll()
        else:
            os.chmod(log, S_IWUSR | S_IREAD)
            logw = open(log, "a+")
            logw.write('Submitted scancode information for scan ' + self.scan.text() + ' to Pharmacy Informatics on ' + str(datetime.datetime.now()) + ' by ' + self.user.text() + '\n')
            sender = '@mskcc.org'
            receiver = ['@mskcc.org','@mskcc.org']
            #receiver = 'molinar1@mskcc.org'
            message = ("Subject: New KBMA Barcode Request\nThis is a test from my Python app\n" +
                        "\nSubmitted by: "+ user+" ("+u+"@mskcc.org)"+
                        "\nLocation: "+loc+
                        "\nNumber: "+str(num)+
                        "\nRaw Scan:"+self.scan.text()+
                        "\nSRx Scan: "+self.sndc.text()+
                        "\nNDC: " + self.ndc.text()+
                        "\nDrug Name: "+self.drug.text()+
                        "\nBrand Name: " + self.bname.text() +
                        "\nStrength: "+self.str.text()+
                        "\nManufacturer: "+self.mfr.text()+
                        "\nPackage Desc: "+self.pack.text()+
                        "\nRoute: "+self.route.text()+
                        "\nForm: "+self.dform.text()+
                        "\nLot: "+self.lot.text()+
                        "\nExp: "+self.exp.text())

            try:
                smtpObj = smtplib.SMTP('mskrelay.mskcc.org', 25)
                smtpObj.sendmail(sender, receiver, message)
                #print("sent")
            except Exception as e:
                logw.write('Error Code 1: Submit not successful! Details: %s\n' % e)
            s = [self.sndc.text(),
                 self.ndc.text(),
                 self.drug.text(),
                 self.bname.text(),
                 self.mfr.text(),
                 self.str.text(),
                 self.dform.text(),
                 self.pack.text(),
                 self.user.text(),
                 self.phone.text(),
                 email]
                 #self.user.text()+'@mskcc.org']
            cnx0 = pyodbc.connect('Driver={SQL Server};'
                            'Server=;'
                           # 'Server=;'
                            'Database=;'
                           # 'Database=;'
                            'uid=;'
                            'pwd=;')
                           # 'pwd=;')

            c0 = cnx0.cursor()
            rows = c0.execute('exec MSKKBMA.KBMASaveDrugBarcodeInfoCollection '
                              '@BarCode = ?, '
                              '@NDC_Num = ?, '
                              '@DrugName = ?, '
                              '@BrandName = ?, '
                              '@Manufacturer = ?, '
                              '@StrConcentration = ?, '
                              '@DosageForm = ?, '
                              '@TotalVolume = ?, '
                              '@EnterBy = ?, '
                              '@Phone_Num = ?, '
                              '@EmailStr = ?;', s)
            c0.commit()
            c0.close()

            QMessageBox.question(self, 'Message', "Submitted to Pharmacy Informatics", QMessageBox.Ok, QMessageBox.Ok)
            self.scan.clear()
            self.scan.setFocus()
            self.drug.setText('')
            self.srx.setText('')
            self.scan.clear()
            self.mfr.clear()
            self.route.clear()
            self.bname.clear()
            self.ndc.clear()
            self.pack.clear()
            self.str.clear()
            self.dform.clear()
            self.lot.clear()
            self.exp.clear()
            self.lexp.clear()
            self.sndc.clear()
            self.scans1.clear()
            logw.close()
            os.chmod(log, S_IREAD | S_IRGRP | S_IROTH)
        self.timer.start(14400000)

    def add2Lexp(self):
        alexp = open('lexp.txt',"a+")
        alexp.write('\n'+self.newLexp.text())
        alexp.close()
        os.chmod(log, S_IWUSR | S_IREAD)
        logw = open(log, "a+")
        logw.write("New Item added: " + self.newLexp.text() + "\n")
        logw.close()
        os.chmod(log, S_IREAD | S_IRGRP | S_IROTH)
        # Updating the table view#
        lines = []
        with open('lexp.txt', 'r') as reader:
            line = reader.readlines()

            for item in line:
                item.replace('\n', '')
                # print(item.rstrip())
                lines.append(item.rstrip())
            # print(lines)
            self.lexpt.setHorizontalHeaderLabels(["SCM Item"])
            self.lexpt.setRowCount(len(lines))

            for line in enumerate(lines):
                # data = line[1]
                for column, data in enumerate(line):
                    # print(line)
                    self.lexpt.setItem(line[0], column, QTableWidgetItem(str(line[1])))
        # This prevents edit from the table itself#

        self.lexpt.resizeColumnsToContents()
        self.lexpt.setEditTriggers(QAbstractItemView.NoEditTriggers)
        QMessageBox.information(self, 'Message', "New Item Added: "+ self.newLexp.text(), QMessageBox.Ok, QMessageBox.Ok)
        self.newLexp.clear()
        self.newLexp.setFocus()
        self.timer.start(14400000)

    def rem2Lexp(self):
        row = self.lexpt.currentRow()
        item = self.lexpt.item(row,0).text()
        #print(item)
        ln = 1

        for line in fileinput.input('lexp.txt', inplace=True):
            if item in line:
                continue
            #print(line, end='')

        #Updating the table view#
        lines = []
        with open('lexp.txt', 'r') as reader:
            line = reader.readlines()

            for item in line:
                item.replace('\n', '')
                # print(item.rstrip())
                lines.append(item.rstrip())
            # print(lines)
            self.lexpt.setHorizontalHeaderLabels(["SCM Item"])
            self.lexpt.setRowCount(len(lines))

            for line in enumerate(lines):
                # data = line[1]
                for column, data in enumerate(line):
                    # print(line)
                    self.lexpt.setItem(line[0], column, QTableWidgetItem(str(line[1])))
        # This prevents edit from the table itself#

        self.lexpt.resizeColumnsToContents()
        self.lexpt.setEditTriggers(QAbstractItemView.NoEditTriggers)
        self.newLexp.setFocus()
        QMessageBox.information(self, 'Message', "Item Removed" + self.newLexp.text(), QMessageBox.Ok, QMessageBox.Ok)
        self.timer.start(14400000)

    def getPi(self):
        self.piTable.clear()
        cnx1 = pyodbc.connect('Driver={SQL Server};'
                              # 'Server=;'
                              'Server=;'
                              'Database=;'
                              # 'Database=;'
                              'UID=;'
                              'PWD=;')
                              # 'PWD=;')

        c1 = cnx1.cursor()
        try:
            c1.execute("SET NOCOUNT ON EXEC KBMAGetPI;")
        except pyodbc.Error as err:
            os.chmod(log, S_IWUSR | S_IREAD)
            logw = open(log, "a+")
            logw.write('Error Code 2: Association Requests not found! Details: %s\n' % err)
            logw.close()
            os.chmod(log, S_IREAD | S_IRGRP | S_IROTH)
        row = c1.fetchall()
        rc1 = len(row)
        #rc1 = c1.rowcount
        #To show all outstanding requests#
        if rc1 == 0:
            #print("No Outstanding Requests.")
            QMessageBox.information(self, 'Message', "No Outstanding Requests.", QMessageBox.Ok, QMessageBox.Ok)
            self.piTable.setRowCount(0)
            self.piTable.setColumnCount(0)
            winsound.PlaySound('done.wav',winsound.SND_ALIAS)
        else:
            self.piTable.setColumnCount(14)
            self.piTable.setRowCount(rc1)
            self.piTable.setHorizontalHeaderLabels(column[0] for column in c1.description)
            #self.piTable.setRowCount(10)
            for row_number, row_data in enumerate(row):
                #self.piTable.insertRow(row_number)
                for column_number, data in enumerate(row_data):
                    self.piTable.setItem(row_number, column_number, QTableWidgetItem(str(data)))
            #self.piTable.show()
            # To prevent editing in window, use this code. This will not, however, commit back to DB#
            self.piTable.setEditTriggers(QAbstractItemView.NoEditTriggers)
            self.piTable.resizeColumnsToContents()
        c1.close()
        self.timer.start(14400000)

    def submitPI(self):
        global piPlayer
        global piSubmit1
        global piSubmit2
        global ws
        global log
        row = self.piTable.currentRow()
        os.chmod(log, S_IWUSR | S_IREAD)
        logw = open(log, 'a+')
        if row <0:
            #print("Nothing Selected")
            QMessageBox.warning(self, 'Message', 'Nothing selected.', QMessageBox.Ok, QMessageBox.Ok)
        else:
            item = self.piTable.item(row, 0).text()
            #print(item)
            pi = (item, ws)
            dtm = datetime.datetime.now()
            cnx1 = pyodbc.connect('Driver={SQL Server};'
                                  'Server=;'
                                  # 'Server=;'
                                  'Database=;'
                                  # 'Database=;'
                                  'uid=;'
                                  'PWD=;')
                                 # 'PWD=;')

            c1 = cnx1.cursor()
            qr = "SET NOCOUNT ON EXEC MSKKBMA.KBMAVerifySRx @barcode = ?, @LocationStr = ?;"
            try:
                c1.execute(qr, pi)
                #c1.commit()
            except pyodbc.Error as err:
                logw.write('Error Code 3: Association Requests not performed! Details: %s\n' % err)
            rp = c1.fetchall()
            #print(rp)
            if not rp:
                QMessageBox.warning(self, 'Warning', 'You have not associated this item in SRx.', QMessageBox.Ok, QMessageBox.Ok)
            elif rp[0][4] is None:
                QMessageBox.warning(self, 'Warning', 'The Scancode is not associated to any Stock Items for KBMA.', QMessageBox.Ok, QMessageBox.Ok)
            else:
                rows = c1.execute("exec MSKKBMA.KBMAUpdateDrugBarcodeInfoCollection @BarCode = ?", item)
                #c1.commit()
                try:
                    c1.execute("SET NOCOUNT ON EXEC KBMAGetPI;")
                except pyodbc.Error as err:
                    logw.write('Error Code 4: Association Requests refresh not successful! Details: %s\n' % err)
                row = c1.fetchall()
                rf = len(row)
                # print(rf)
                if rf == 0:
                    # print("No Outstanding Requests.")
                    QMessageBox.question(self, 'Message', "No Outstanding Requests.", QMessageBox.Ok,QMessageBox.Ok)
                    self.piTable.setRowCount(0)
                    self.piTable.setColumnCount(0)
                    winsound.PlaySound('done.wav', winsound.SND_ALIAS)
                else:
                    self.piTable.setColumnCount(14)
                    self.piTable.setRowCount(rf)
                    self.piTable.setHorizontalHeaderLabels(column[0] for column in c1.description)
                    # self.piTable.setRowCount(10)
                    for row_number, row_data in enumerate(row):
                        # self.piTable.insertRow(row_number)
                        for column_number, data in enumerate(row_data):
                            self.piTable.setItem(row_number, column_number, QTableWidgetItem(str(data)))
                    # self.piTable.show()
                    # To prevent editing in window, use this code. This will not, however, commit back to DB#
                    self.piTable.setEditTriggers(QAbstractItemView.NoEditTriggers)
                    self.piTable.resizeColumnsToContents()

                    winsound.PlaySound('another.wav', winsound.SND_ALIAS)
                    winsound.PlaySound('airhorn.wav', winsound.SND_ALIAS)
                    os.chmod(log, S_IWUSR | S_IREAD)
                    logw = open(log, "a+")
                    logw.write('Request for Scan ' + item + ' completed on ' + str(dtm) + " by " + self.user.text() + ".\n")
                    logw.close()
                    os.chmod(log, S_IREAD | S_IRGRP | S_IROTH)
            c1.commit()
            c1.close()
        self.timer.start(14400000)
        logw.close()
        os.chmod(log, S_IREAD | S_IRGRP | S_IROTH)

    def parseScan(self):
        global log
        global ws
        # print(ws)
        nm = QPixmap("NoMatch.png").scaledToHeight(25)
        m = QPixmap("Match.png").scaledToHeight(25)
        os.chmod(log, S_IWUSR | S_IREAD)
        logw = open(log, 'a+')
        if len(self.drug.text()) != 0 or len(self.mfr.text()) != 0 \
                or len(self.bname.text()) != 0 or len(self.str.text()) != 0 \
                or len(self.dform.text()) != 0 or len(self.pack.text()) != 0:
            quest1 = QMessageBox.question(self, "Warning","You're about to restart the Receiving process.\nAre you sure you want to proceed?",QMessageBox.Yes | QMessageBox.No, QMessageBox.Yes)
            if quest1 == QMessageBox.Yes:

                logw.write('Receiving Information for this product was not submitted. The Process was restarted.\n')
                self.ndc.clear()
                self.srx.clear()
                self.mfr.clear()
                self.drug.clear()
                self.route.clear()
                self.bname.clear()
                self.scans1.clear()
                self.sndc.clear()
                self.pack.clear()
                self.str.clear()
                self.lot.clear()
                self.exp.clear()
                self.dform.clear()
                self.image_res.clear()
                self.submitR.setEnabled(False)
                self.search1.setEnabled(False)
            else:
                self.scan.undo()
            self.scan.setFocus()
        else:
            self.mfr.clear()
            self.drug.clear()
            self.route.clear()
            self.bname.clear()
            self.ndc.clear()
            self.pack.clear()
            self.str.clear()
            self.dform.clear()
            self.image_res.clear()
            self.lot.clear()
            self.exp.clear()
            scan = self.scan.text()
            scan = str(re.sub(r'[^\w]', '', scan))
            try:
                bc_10 = (r'(?P<ndc>[0-9]{10})?$')
                bc_10 = re.compile(bc_10)

                bc_10_1 = (r'0(?P<ndc>[0-9]{10})?$')
                bc_10_1 = re.compile(bc_10_1)

                bc_11 = (r'(?P<ndc>[0-9]{11})?$')
                bc_11 = re.compile(bc_11)

                bc_exp_lot = (
                    r'(01|\(01\))(?P<pad>[0-9]{3})(?P<ndc>[0-9]{10})17(?P<exp>[0-9]{6})10(?P<lot>[\x21-\x22\x25-\x2F\x30-\x39\x41-\x5A\x5F\x61-\x7A]{0,20})?$')
                bc_exp_lot = re.compile(bc_exp_lot)

                bc_cd_exp_lot = (
                    r'(01|\(01\))(?P<pad>[0-9]{3})(?P<ndc>[0-9]{10})(?P<check_dig>[0-9]{1})17(?P<exp>[0-9]{6})10(?P<lot>[\x21-\x22\x25-\x2F\x30-\x39\x41-\x5A\x5F\x61-\x7A]{0,20})?$')
                bc_cd_exp_lot = re.compile(bc_cd_exp_lot)

                bc_cd_lot_exp = (
                    r'(01|\(01\))(?P<pad>[0-9]{3})(?P<ndc>[0-9]{10})(?P<check_dig>[0-9]{1})10(?P<lot>([\x21-\x22\x25-\x2F\x30-\x39\x3A-\x3F\x41-\x5A\x5F\x61-\x7A]{0,20}))17(?P<exp>[0-9]{6})?$')
                bc_cd_lot_exp = re.compile(bc_cd_lot_exp)

                bc_lot_exp = (
                    r'(01|\(01\))(?P<pad>[0-9]{3})(?P<ndc>[0-9]{10})10(?P<lot>[\x21-\x22\x25-\x2F\x30-\x39\x41-\x5A\x5F\x61-\x7A]{0,20})17(?P<exp>[0-9]{6})?$')
                bc_lot_exp = re.compile(bc_lot_exp)

                bc_serial_lot_exp = (
                    r'(01|\(01\))(?P<pad>[0-9]{3})(?P<ndc>[0-9]{10})(?P<check_dig>[0-9]{1})21(?P<serial>[0-9,A-Z]{10,14})(17(?P<exp>[0-9]{6})10(?P<lot>[\x21-\x22\x25-\x2F\x30-\x39\x41-\x5A\x5F\x61-\x7A]{0,20}))?$')
                bc_serial_lot_exp = re.compile(bc_serial_lot_exp)

                bc_exp_lot_serial = (
                    r'(01|\(01\))(?P<pad>[0-9]{3})(?P<ndc>[0-9]{10})(?P<check_dig>[0-9]{1})17(?P<exp>[0-9]{6})10(?P<lot>[\x21-\x22\x25-\x2F\x30-\x39\x41-\x5A\x5F\x61-\x7A]{0,20})21(?P<serial>[0-9,A-Z]{10,14})?$')
                bc_exp_lot_serial = re.compile(bc_exp_lot_serial)

                bc_lot_exp_serial = (
                    r'(01|\(01\))(?P<pad>[0-9]{3})(?P<ndc>[0-9]{10})(?P<check_dig>[0-9]{1})10(?P<lot>[\x21-\x22\x25-\x2F\x30-\x39\x41-\x5A\x5F\x61-\x7A]{0,20})17(?P<exp>[0-9]{6})21(?P<serial>[0-9,A-Z]{10,14})?$')
                bc_lot_exp_serial = re.compile(bc_lot_exp_serial)

                bc_plain = (r'(01|\(01\))(?P<pad>[0-9]{3})(?P<ndc>[0-9]{10})?$')
                bc_plain = re.compile(bc_plain)

                bc_plain1 = (r'(01|\(01\))(?P<pad>[0-9]{3})(?P<ndc>[0-9]{10})[0-9]*?$')
                bc_plain1 = re.compile(bc_plain1)

                bc_lot = (
                    r'(01|\(01\))(?P<pad>[0-9]{3})(?P<ndc>[0-9]{10})(?P<check_dig>[0-9]{1})10(?P<lot>[\x21-\x22\x25-\x2F\x30-\x39\x41-\x5A\x5F\x61-\x7A]{0,20})?$')
                bc_lot = re.compile(bc_lot)

                bc_repack = (r'3(?P<ndc>[0-9]{10})?$')
                bc_repack = re.compile(bc_repack)

                bc_repack1 = (r'3(?P<ndc>[0-9]{10})(?P<check_dig>[0-9]{1})?$')
                bc_repack1 = re.compile(bc_repack1)

                bc_repack_exp = (r'3(?P<ndc>[0-9]{11})(?P<check_dig>[0-9]{1})(?P<exp>[0-9]{4})?$')
                bc_repack_exp = re.compile(bc_repack_exp)

                bc_repack_exp1 = (r'3(?P<ndc>[0-9]{10})(?P<check_dig>[0-9]{1})(?P<exp>[0-9]{4})?$')
                bc_repack_exp1 = re.compile(bc_repack_exp1)

                lot_exp_bc = (r'17(?P<exp>[0-9]{6})10(?P<lot>[\x21-\x22\x25-\x2F\x30-\x39\x41-\x5A\x5F\x61-\x7A]{0,20})?$')
                lot_exp_bc = re.compile(lot_exp_bc)

                bc_otc = (r'003(?P<ndc>[0-9]{9,10})(?P<check_dig>[0-9]{1})?$')
                bc_otc = re.compile(bc_otc)

                if not (scan.startswith('01') or scan.startswith('(01)') or scan.startswith('00')) and (scan.startswith('17') and len(scan) > 10):
                    if self.sndc.text() != '':
                        match = lot_exp_bc.match(scan)
                        if 'lot' in match.groupdict():
                            lot = match.group('lot')
                            # print('Lot:', match.group('lot'))
                        else:
                            lot = ''
                        if 'exp' in match.groupdict():
                            if len(match.group('exp')) == 6:
                                if match.group('exp').endswith('00'):
                                    exp = match.group('exp')[2:4] + "/" + match.group('exp')[0:2]
                                    # print('Exp:', match.group('exp')[2:4] + "/" + match.group('exp')[0:2])
                                else:
                                    exp = match.group('exp')[2:4] + "/" + match.group('exp')[4:6] + "/" + match.group(
                                        'exp')[0:2]
                                    # print('Exp:',match.group('exp')[2:4] + "/" + match.group('exp')[4:6] + "/" + match.group('exp')[0:2])
                            else:
                                exp = match.group('exp')[0:2] + "/" + match.group('exp')[2:4]
                                # print('Exp:', match.group('exp')[0:2] + "/" + match.group('exp')[2:4])
                        else:
                            exp = ''
                    else:
                        QMessageBox.warning(self, 'Warning', "You must first scan the NDC barcode.", QMessageBox.Ok,QMessageBox.Ok)
                        lot = ''
                        exp = ''

                elif scan.startswith("M"):
                    QMessageBox.warning(self, "Warning", "Please use the Dispensing Tab for scanning labels",QMessageBox.Ok, QMessageBox.Ok)
                    self.scan.clear()
                    self.scan.setFocus()
                    lot = ''
                    exp = ''

                else:
                    if scan.startswith('01') or scan.startswith('(01)') or (scan.startswith('00') and len(scan) < 11) or (scan.startswith('17') and len(scan) < 11):
                        match = bc_10.match(scan)
                        if not match:
                            match = bc_11.match(scan)
                            if not match:
                                match = bc_lot_exp.match(scan)
                                if not match:
                                    match = bc_serial_lot_exp.match(scan)
                                    if not match:
                                        match = bc_lot_exp_serial.match(scan)
                                        if not match:
                                            match = bc_exp_lot_serial.match(scan)
                                            if not match:
                                                match = bc_exp_lot.match(scan)
                                                if not match:
                                                    match = bc_plain.match(scan)
                                                    if not match:
                                                        match = bc_plain1.match(scan)
                                                        if not match:
                                                            match = bc_cd_lot_exp.match(scan)
                                                            if not match:
                                                                match = bc_cd_exp_lot.match(scan)
                                                                if not match:
                                                                    match = bc_lot.match(scan)

                    elif scan.startswith('00') and len(scan) > 10:
                        match = bc_10_1.match(scan)
                        if not match:
                            match = bc_otc.match(scan)

                    elif scan.startswith('3') and len(scan) > 10:
                        match = bc_repack.match(scan)
                        if not match:
                            match = bc_repack1.match(scan)
                            if not match:
                                match = bc_repack_exp.match(scan)
                                if not match:
                                    match = bc_repack_exp1.match(scan)

                    elif not (scan.startswith('01') or scan.startswith('(01)') or scan.startswith('00') or (
                            scan.startswith('17') and len(scan) < 11)):
                        match = bc_10.match(scan)
                        if not match:
                            match = bc_11.match(scan)
                            if not match:
                                match = bc_10_1.match(scan)
                                if not match:
                                    match = lot_exp_bc.match(scan)
                    if not match:
                        QMessageBox.warning(self, 'Warning','This is an invalid barcode.\nPlease try scanning another barcode.',QMessageBox.Ok, QMessageBox.Ok)
                        logw.write("Scancode " + self.scan.text() + " is not a valid scan.\n")
                        self.scan.clear()
                        self.srx.clear()
                        self.scan.setFocus()
                        ndc = ''
                        lot = ''
                        exp = ''
                    else:
                        if 'ndc' in match.groupdict():
                            ndc = match.group('ndc')
                            self.scans1.setText(ndc)
                            if scan.startswith('3') and len(scan) > 10:
                                self.sndc.setText('3' + ndc)
                            else:
                                self.sndc.setText(ndc)
                            # print('NDC:', match.group('ndc'))
                        else:
                            pass
                        if 'lot' in match.groupdict():
                            lot = match.group('lot')
                            # print('Lot:', match.group('lot'))
                        else:
                            lot = ''
                        if 'exp' in match.groupdict():
                            if len(match.group('exp')) == 6:
                                if match.group('exp').endswith('00'):
                                    exp = match.group('exp')[2:4] + "/" + match.group('exp')[0:2]
                                    # print('Exp:', match.group('exp')[2:4] + "/" + match.group('exp')[0:2])
                                else:
                                    exp = match.group('exp')[2:4] + "/" + match.group('exp')[4:6] + "/" + match.group('exp')[0:2]
                                    # print('Exp:',match.group('exp')[2:4] + "/" + match.group('exp')[4:6] + "/" + match.group('exp')[0:2])
                            else:
                                exp = match.group('exp')[0:2] + "/" + match.group('exp')[2:4]
                                # print('Exp:', match.group('exp')[0:2] + "/" + match.group('exp')[2:4])
                        else:
                            exp = ''

                        if 'serial' in match.groupdict():
                            serial = match.group('serial')
                            # print('Serial:', match.group('serial'))
                        else:
                            serial = ''

                        cnx = pyodbc.connect('Driver={SQL Server};'
                                             'Server=;'
                                             # 'Server=;'
                                             'Database=;'
                                             # 'Database=;'
                                             'uid=;'
                                             'pwd=;')
                                             # 'pwd=;')

                        cursor = cnx.cursor()
                        qb = "SET NOCOUNT ON EXEC MSKKBMA.KBMAVerifySRx @barcode = ?, @LocationStr = ?;"
                        pb = (self.sndc.text(), ws)
                        # print(pb)
                        try:
                            cursor.execute(qb, pb)
                        except pyodbc.Error as err:
                            logw.write('Error Code 5: ParseScan VerifySRx failed! Details: %s\n' % err)
                        rc = cursor.fetchall()
                        # print(rc)
                        # setting up results #
                        if not rc:
                            qr = "SET NOCOUNT ON EXEC MSKKBMA.KBMAVerifyBarcode @Barcode = ?;"
                            try:
                                cursor.execute(qr, self.sndc.text())
                            except pyodbc.Error as err:
                                logw.write('Error Code 6: ParseScan VerifyBarcode failed! Details: %s\n' % err)
                            rb = cursor.fetchall()
                            # print(rb[0][0])
                            if rb[0][0] == 0:
                                logw.write("Scancode " + scan + " for NDC " + ndc + " not in SRx; New Request needed.\n")
                                self.srx.setText("Not in SRx")
                                winsound.PlaySound('buzzer.wav', winsound.SND_ALIAS)
                                self.image_res.setPixmap(nm)
                                self.image_res.setMaximumWidth(25)
                                self.drug.setText('')
                                self.mfr.clear()
                                self.route.clear()
                                self.bname.clear()
                                self.ndc.clear()
                                self.pack.clear()
                                self.str.clear()
                                self.dform.clear()
                                self.search1.setEnabled(True)
                                self.submitR.setEnabled(False)
                            else:
                                logw.write("A Duplicate Barcode Request for Scancode " + scan + " (NDC " + ndc + ") exists.\n")
                                winsound.PlaySound('SystemExclamation', winsound.SND_ALIAS)
                                QMessageBox.warning(self, 'Duplicate Request',"This Barcode has already been submitted.\nNo further action required.",QMessageBox.Ok, QMessageBox.Ok)
                                self.scan.clear()
                                self.sndc.clear()
                                self.srx.clear()
                                self.scan.setFocus()
                                self.search1.setEnabled(False)
                                self.submitR.setEnabled(False)
                        else:
                            # print(rc[0][1])
                            logw.write("Match Found for Scancode " + scan + " (NDC " + ndc + "): " + rc[0][1] + " by " + rc[0][3] + "\n")
                            self.srx.setText(rc[0][1] + " by " + rc[0][3])
                            winsound.PlaySound('bing.wav', winsound.SND_ALIAS)
                            self.image_res.setPixmap(m)
                            self.image_res.setMaximumWidth(25)
                            self.search1.setEnabled(False)
                            self.submitR.setEnabled(False)

                        # self.srx.clear()
                        cursor.commit()
                        cursor.close()
                self.lot.setText(lot)
                self.exp.setText(exp)
            except Exception as e:
                logw.write('Error Code 7: ParseScan failed! Details: %s\n'% e)
        logw.close()
        os.chmod(log, S_IREAD | S_IRGRP | S_IROTH)
        self.scan.selectAll()
        self.timer.start(14400000)

    def parseScan2(self):
        global log
        global passScan
        global failScan
        global ws

        self.image_res2.clear()
        scan = self.scan2.text()
        nm = QPixmap("NoMatch.png").scaledToHeight(45)
        m = QPixmap("Match.png").scaledToHeight(45)
        cnx2 = pyodbc.connect('Driver={SQL Server};'
                              # 'Server=;'
                              'Server=;'
                              # 'Database=;'
                              'Database=;'
                              'UID=;'
                              'PWD=;')
                             # 'PWD=;')
        os.chmod(log, S_IWUSR | S_IREAD)
        logw = open(log, "a+")
        try:

            if scan.startswith('M'):
                # print('2')
                if (len(self.res.text()) == 0 or len(self.lexpi.text()) == 0):
                    self.ndc2.clear()
                    self.exp2.clear()
                    self.lot2.clear()
                    self.lastScan.clear()
                    self.srp2.clear()
                    # print(self.f.count())

                    if len(scan) == 10:
                        mNum = scan
                        self.scan3.setText(mNum)
                        # scan3 = mNum
                        doseNum = None
                        self.doseN.setText(doseNum)
                    elif len(scan) > 10:
                        mNum = scan[0:10]
                        doseNum = scan[scan.find(' ') + 1:len(scan)]
                        self.doseN.setText(doseNum)
                        self.scan3.setText(mNum)
                    res = []

                    try:
                        qd = "SET NOCOUNT ON exec MSKKBMA.KBMADisp @MMNum = ?;"
                        cursor2 = cnx2.cursor()
                        cursor2.execute(qd, mNum)
                    except pyodbc.Error as err:
                        logw.write('Error Code 8: ParseScan2 OrderLookup failed! Details: %s\n' % err)
                    row = cursor2.fetchall()
                    # print(row)
                    if not row:
                        QMessageBox.warning(self, 'Warning', 'The Order scanned is no longer valid.', QMessageBox.Ok, QMessageBox.Ok)
                        self.scan2.clear()
                        self.srx2.clear()
                        self.scan2.setFocus()
                    else:

                        for item in row:
                            # print(item[2])
                            if item[1] is None:
                                orderName = item[0]
                                self.srx2.setText(item[0])
                            else:
                                orderName = item[0] + ' ' + item[1]
                                self.srx2.setText(item[0] + ' ' + item[1])
                            res.append(str(item[2]))
                            self.res.setText(', '.join(res))

                            with open('lexp.txt', 'r') as lexp:
                                lexpl = lexp.readlines()
                                for line in lexpl:
                                    if item[0] == line.rstrip():
                                        self.lexpi.setText("Y")
                                        self.submitD.setHidden(False)
                                        self.addItem.setHidden(False)
                                        break
                                    else:
                                        self.lexpi.clear()
                                        self.submitD.setHidden(True)
                                        self.addItem.setHidden(True)
                        # print(self.res.text())
                        # print(self.srx2.text())
                        cursor2.close()
                        logw.write('\n' + "Verifying Order " + orderName + "\n")
                    self.scan2.clear()
                    self.scan2.setFocus()
                    # mNum =''

                elif len(self.lexpi.text()) != 0:
                    question = QMessageBox.question(self, 'Message',"Are you sure you want to cancel this dispense action?",QMessageBox.Yes | QMessageBox.No, QMessageBox.Yes)
                    if question == QMessageBox.Yes:
                        logw.write('Dispensing Information for this order was not submitted.\n')
                        self.ndc2.clear()
                        self.lot2.clear()
                        self.exp2.clear()
                        self.image_res2.clear()
                        self.scan3.clear()
                        mNum = ''
                        item = 16
                        while item in range(self.f.count()):
                            # item += 1
                            self.it = self.f.itemAt(item).widget()
                            self.f.removeWidget(self.it)
                            self.it.deleteLater()
                            self.ndcLine = 4
                            del self.it
                            # item += 1

                        # print(self.f.count())
                        self.res.clear()
                        self.scan2.clear()
                        self.srx2.clear()
                        self.srp2.clear()
                        self.scan2.setFocus()
                        self.lexpi.clear()
                        res = []
                        self.scan2.setFocus()
                    else:
                        self.scan2.undo()
                        self.scan2.undo()
                        self.srx2.clear()
                        self.srp2.clear()
                        self.scan2.setFocus()
                    self.lastScan.clear()

            # If no M-number has been scanned#
            elif not scan.startswith('M'):
                scan = str(re.sub(r'[^\w]', '', scan))

                bc_10 = (r'(?P<ndc>[0-9]{10})?$')
                bc_10 = re.compile(bc_10)

                bc_10_1 = (r'0(?P<ndc>[0-9]{10})?$')
                bc_10_1 = re.compile(bc_10_1)

                bc_11 = (r'(?P<ndc>[0-9]{11})?$')
                bc_11 = re.compile(bc_11)

                bc_exp_lot = (
                    r'(01|\(01\))(?P<pad>[0-9]{3})(?P<ndc>[0-9]{10})17(?P<exp>[0-9]{6})10(?P<lot>[\x21-\x22\x25-\x2F\x30-\x39\x41-\x5A\x5F\x61-\x7A]{0,20})?$')
                bc_exp_lot = re.compile(bc_exp_lot)

                bc_cd_exp_lot = (
                    r'(01|\(01\))(?P<pad>[0-9]{3})(?P<ndc>[0-9]{10})(?P<check_dig>[0-9]{1})17(?P<exp>[0-9]{6})10(?P<lot>[\x21-\x22\x25-\x2F\x30-\x39\x41-\x5A\x5F\x61-\x7A]{0,20})?$')
                bc_cd_exp_lot = re.compile(bc_cd_exp_lot)

                bc_cd_lot_exp = (
                    r'(01|\(01\))(?P<pad>[0-9]{3})(?P<ndc>[0-9]{10})(?P<check_dig>[0-9]{1})10(?P<lot>([\x21-\x22\x25-\x2F\x30-\x39\x3A-\x3F\x41-\x5A\x5F\x61-\x7A]{0,20}))17(?P<exp>[0-9]{6})?$')
                bc_cd_lot_exp = re.compile(bc_cd_lot_exp)

                bc_lot_exp = (
                    r'(01|\(01\))(?P<pad>[0-9]{3})(?P<ndc>[0-9]{10})10(?P<lot>[\x21-\x22\x25-\x2F\x30-\x39\x41-\x5A\x5F\x61-\x7A]{0,20})17(?P<exp>[0-9]{6})?$')
                bc_lot_exp = re.compile(bc_lot_exp)

                bc_serial_lot_exp = (
                    r'(01|\(01\))(?P<pad>[0-9]{3})(?P<ndc>[0-9]{10})(?P<check_dig>[0-9]{1})21(?P<serial>[0-9,A-Z]{10,14})(17(?P<exp>[0-9]{6})10(?P<lot>[\x21-\x22\x25-\x2F\x30-\x39\x41-\x5A\x5F\x61-\x7A]{0,20}))?$')
                bc_serial_lot_exp = re.compile(bc_serial_lot_exp)

                bc_exp_lot_serial = (
                    r'(01|\(01\))(?P<pad>[0-9]{3})(?P<ndc>[0-9]{10})(?P<check_dig>[0-9]{1})17(?P<exp>[0-9]{6})10(?P<lot>[\x21-\x22\x25-\x2F\x30-\x39\x41-\x5A\x5F\x61-\x7A]{0,20})21(?P<serial>[0-9,A-Z]{10,14})?$')
                bc_exp_lot_serial = re.compile(bc_exp_lot_serial)

                bc_lot_exp_serial = (
                    r'(01|\(01\))(?P<pad>[0-9]{3})(?P<ndc>[0-9]{10})(?P<check_dig>[0-9]{1})10(?P<lot>[\x21-\x22\x25-\x2F\x30-\x39\x41-\x5A\x5F\x61-\x7A]{0,20})17(?P<exp>[0-9]{6})21(?P<serial>[0-9,A-Z]{10,14})?$')
                bc_lot_exp_serial = re.compile(bc_lot_exp_serial)

                bc_plain = (r'(01|\(01\))(?P<pad>[0-9]{3})(?P<ndc>[0-9]{10})?$')
                bc_plain = re.compile(bc_plain)

                bc_plain1 = (r'(01|\(01\))(?P<pad>[0-9]{3})(?P<ndc>[0-9]{10})[0-9]*?$')
                bc_plain1 = re.compile(bc_plain1)

                bc_lot = (
                    r'(01|\(01\))(?P<pad>[0-9]{3})(?P<ndc>[0-9]{10})(?P<check_dig>[0-9]{1})10(?P<lot>[\x21-\x22\x25-\x2F\x30-\x39\x41-\x5A\x5F\x61-\x7A]{0,20})?$')
                bc_lot = re.compile(bc_lot)

                bc_repack = (r'3(?P<ndc>[0-9]{10})?$')
                bc_repack = re.compile(bc_repack)

                bc_repack1 = (r'3(?P<ndc>[0-9]{10})(?P<check_dig>[0-9]{1})?$')
                bc_repack1 = re.compile(bc_repack1)

                bc_repack_exp = (r'3(?P<ndc>[0-9]{11})(?P<check_dig>[0-9]{1})(?P<exp>[0-9]{4})?$')
                bc_repack_exp = re.compile(bc_repack_exp)

                bc_repack_exp1 = (r'3(?P<ndc>[0-9]{10})(?P<check_dig>[0-9]{1})(?P<exp>[0-9]{4})?$')
                bc_repack_exp1 = re.compile(bc_repack_exp1)

                lot_exp_bc = (r'17(?P<exp>[0-9]{6})10(?P<lot>[\x21-\x22\x25-\x2F\x30-\x39\x41-\x5A\x5F\x61-\x7A]{0,20})?$')
                lot_exp_bc = re.compile(lot_exp_bc)

                bc_otc = (r'003(?P<ndc>[0-9]{9,10})(?P<check_dig>[0-9]{1})?$')
                bc_otc = re.compile(bc_otc)

                if (scan.startswith('17') and len(scan) > 10) and not (scan.startswith('01') or scan.startswith('(01)') or scan.startswith('00')):
                    #print('nondrugscan')
                    if len(self.res.text()) == 0:
                        # print('here')
                        QMessageBox.question(self, 'Message', "You must scan the order label first.",
                                             QMessageBox.Ok,
                                             QMessageBox.Ok)
                        self.scan2.setFocus()
                        self.scan2.selectAll()
                        lot = ''
                        exp = ''


                    elif (self.f.count() == 16 and len(self.ndc2.text()) == 0):
                        QMessageBox.question(self, 'Message', "Scan NDC Barcode First", QMessageBox.Ok, QMessageBox.Ok)
                        self.scan2.setFocus()
                        self.scan2.selectAll()
                        lot = ''
                        exp = ''


                    elif self.f.count() == 22 and (len(self.ndc2.text()) == 0 or len(self.f.itemAt(17).widget().text()) == 0):
                        QMessageBox.question(self, 'Message', "Scan NDC Barcode First", QMessageBox.Ok, QMessageBox.Ok)
                        self.scan2.setFocus()
                        self.scan2.selectAll()
                        lot = ''
                        exp = ''


                    elif self.f.count() > 22 and (len(self.ndc2.text()) == 0 or len(self.f.itemAt(17).widget().text()) == 0 or len(self.f.itemAt(23).widget().text()) == 0):
                        QMessageBox.question(self, 'Message', "Scan NDC Barcode First", QMessageBox.Ok, QMessageBox.Ok)
                        self.scan2.setFocus()
                        self.scan2.selectAll()
                        lot = ''
                        exp = ''

                    else:
                        match = lot_exp_bc.match(scan)
                        # print(match.groupdict())
                        if 'lot' in match.groupdict():
                            lot = match.group('lot')
                            # print('Lot:', match.group('lot'))
                        else:
                            lot = ''
                        if 'exp' in match.groupdict():
                            if len(match.group('exp')) == 6:
                                if match.group('exp').endswith('00'):
                                    exp = match.group('exp')[2:4] + "/" + match.group('exp')[0:2]
                                    # print('Exp:', match.group('exp')[2:4] + "/" + match.group('exp')[0:2])
                                else:
                                    exp = match.group('exp')[2:4] + "/" + match.group('exp')[
                                                                          4:6] + "/" + match.group(
                                        'exp')[0:2]
                                    # print('Exp:',match.group('exp')[2:4] + "/" + match.group('exp')[4:6] + "/" + match.group('exp')[0:2])
                            else:
                                exp = match.group('exp')[0:2] + "/" + match.group('exp')[2:4]
                                # print('Exp:', match.group('exp')[0:2] + "/" + match.group('exp')[2:4])
                        else:
                            exp = ''

                        if self.f.count() == 16:
                            self.exp2.setText(exp)
                            self.lot2.setText(lot)
                        elif self.f.count() == 22:
                            ex = self.f.itemAt(21)
                            lx = self.f.itemAt(19)
                            exw = ex.widget()
                            exw.setText(exp)
                            lxw = lx.widget()
                            lxw.setText(lot)
                        elif self.f.count() > 22:
                            ex = self.f.itemAt(27)
                            lx = self.f.itemAt(25)
                            exw = ex.widget()
                            exw.setText(exp)
                            lxw = lx.widget()
                            lxw.setText(lot)


                else:
                    # print('drugscan')
                    try:
                        if scan.startswith('01') or scan.startswith('(01)') or (scan.startswith('00') and len(scan) < 11) or (scan.startswith('17') and len(scan) < 11):
                            match = bc_10.match(scan)
                            if not match:
                                match = bc_11.match(scan)
                                if not match:
                                    match = bc_lot_exp.match(scan)
                                    if not match:
                                        match = bc_serial_lot_exp.match(scan)
                                        if not match:
                                            match = bc_lot_exp_serial.match(scan)
                                            if not match:
                                                match = bc_exp_lot_serial.match(scan)
                                                if not match:
                                                    match = bc_exp_lot.match(scan)
                                                    if not match:
                                                        match = bc_plain.match(scan)
                                                        if not match:
                                                            match = bc_plain1.match(scan)
                                                            if not match:
                                                                match = bc_cd_lot_exp.match(scan)
                                                                if not match:
                                                                    match = bc_cd_exp_lot.match(scan)
                                                                    if not match:
                                                                        match = bc_lot.match(scan)

                        elif scan.startswith('00') and len(scan) > 10:
                            match = bc_10_1.match(scan)
                            if not match:
                                match = bc_otc.match(scan)

                        elif scan.startswith('3') and len(scan) > 10:
                            match = bc_repack.match(scan)
                            if not match:
                                match = bc_repack1.match(scan)
                                if not match:
                                    match = bc_repack_exp.match(scan)
                                    if not match:
                                        match = bc_repack_exp1.match(scan)

                        elif not (scan.startswith('01') or scan.startswith('(01)') or scan.startswith('00') or (scan.startswith('17') and len(scan) < 11)):
                            match = bc_10.match(scan)
                            if not match:
                                match = bc_11.match(scan)
                                if not match:
                                    match = bc_10_1.match(scan)
                                    if not match:
                                        match = lot_exp_bc.match(scan)

                        if not match:
                            QMessageBox.warning(self, 'Warning','This is an invalid barcode.\nPlease try scanning another barcode.',QMessageBox.Ok, QMessageBox.Ok)
                            self.scan.clear()
                            self.srp2.setText('Last Scan was an Invalid Barcode')
                            self.scan.setFocus()
                            lot = ''
                            exp = ''
                        else:
                            if 'ndc' in match.groupdict():
                                ndc1 = match.group('ndc')
                                self.sndc1 = QLineEdit()
                                if scan.startswith('3') and len(scan) > 10:
                                    self.sndc1.setText('3' + ndc1)
                                else:
                                    self.sndc1.setText(ndc1)

                            else:
                                pass
                            # print('NDC:', ndc)
                            if 'lot' in match.groupdict():
                                lot = match.group('lot')

                            else:
                                lot = ''
                            # print('Lot:', lot)
                            if 'exp' in match.groupdict():
                                if len(match.group('exp')) == 6:
                                    if match.group('exp').endswith('00'):
                                        exp = match.group('exp')[2:4] + "/" + match.group('exp')[0:2]
                                        # print('Exp:', match.group('exp')[2:4] + "/" + match.group('exp')[0:2])
                                    else:
                                        exp = match.group('exp')[2:4] + "/" + match.group('exp')[4:6] + "/" + match.group(
                                            'exp')[0:2]
                                        # print('Exp:',match.group('exp')[2:4] + "/" + match.group('exp')[4:6] + "/" + match.group('exp')[0:2])
                                else:
                                    exp = match.group('exp')[0:2] + "/" + match.group('exp')[2:4]
                                    # print('Exp:', match.group('exp')[0:2] + "/" + match.group('exp')[2:4])
                            else:
                                exp = ''
                            # print('Exp:', exp)
                            if 'serial' in match.groupdict():
                                serial = match.group('serial')
                                # print('Serial:', match.group('serial'))
                            else:
                                serial = ''
                            # print('Serial:', serial)
                            # log = open('log.txt', 'a+')
                            if ndc1 in self.res.text():
                                # print('found')
                                logw.write('Correct Scan ' + scan + "\n")
                                winsound.PlaySound('bing.wav', winsound.SND_ALIAS)
                                self.image_res2.setPixmap(m)
                                self.image_res2.setMaximumWidth(45)
                                self.lastScan.setText("P")
                                if self.f.count() == 16:
                                    self.ndc2.setText(ndc1)
                                    self.lot2.setText(lot)
                                    self.exp2.setText(exp)
                                elif self.f.count() == 22:
                                    n = self.f.itemAt(17).widget()
                                    n1 = self.f.itemAt(19).widget()
                                    n2 = self.f.itemAt(21).widget()
                                    n.setText(ndc1)
                                    n1.setText(lot)
                                    n2.setText(exp)
                                elif self.f.count() > 22:
                                    n3 = self.f.itemAt(23).widget()
                                    n4 = self.f.itemAt(25).widget()
                                    n5 = self.f.itemAt(27).widget()
                                    n3.setText(ndc1)
                                    n4.setText(lot)
                                    n5.setText(exp)
                            else:
                                winsound.PlaySound('buzzer.wav', winsound.SND_ALIAS)
                                self.image_res2.setPixmap(nm)
                                self.image_res2.setMaximumWidth(45)
                                self.lastScan.setText("F")
                                logw.write('Incorrect Scan ' + ndc1 + '\n')

                            dl = (datetime.datetime.now(), self.scan3.text(), self.doseN.text(), str(ndc1), lot, exp,self.user.text(), self.lastScan.text())
                            # print(dl)
                            try:
                                q2 = "exec KBMADispR @createdwhen = ?, @MMNum = ?, @dose = ?, @NDC = ?, @Lot = ?, @Exp = ?, @User = ?, @ScanStatus = ?;"
                                cn2 = cnx2.cursor()
                                cn2.execute(q2, dl)
                                cn2.commit()
                                qb = "SET NOCOUNT ON EXEC MSKKBMA.KBMAVerifySRx @barcode = ?, @LocationStr = ?;"
                                pb = (self.sndc1.text(), ws)
                                cn2.execute(qb, pb)
                            except pyodbc.Error as err:
                                # print(err)
                                logw.write('Error Code 9: ParseScan2 VerifySRx failed! Details: %s\n' % err)
                            rc6 = cn2.fetchall()
                            # print(rc6)
                            if not rc6:
                                QMessageBox.warning(self, 'Warning',
                                                    'This Scancode is not associated to any Products in SRx.\nPlease use the Receiving Tab to send information to PI group.',
                                                    QMessageBox.Ok, QMessageBox.Ok)
                                logw.write('Scan not in SRx ' + ndc1 + ' \n')
                                
                                self.srp2.setText("Last Scanned Barcode not in SRx")

                            else:
                                self.srp2.setText(rc6[0][1] + " by " + rc6[0][3])
                                logw.write(rc6[0][1] + " by " + rc6[0][3] + '\n')
                                
                            cn2.close()
                    except Exception:
                        logw.write('Error Code 10: ParseScan2 parsing failed! Details: %s\n'+ traceback.format_exc())
                        


            else:
                QMessageBox.warning(self, 'Warning', 'This is not a valid scan.', QMessageBox.Ok, QMessageBox.Ok)
                self.scan2.selectAll()
                self.scan2.setFocus()

        except Exception:
            logw.write('Error Code 11: ParseScan2 failed! Details: %s\n', traceback.format_exc())
        logw.close()
        os.chmod(log, S_IREAD | S_IRGRP | S_IROTH)
        self.scan2.selectAll()
        self.scan2.setFocus()
        self.timer.start(14400000)

    def parseScan3(self):
        psnf = QPixmap("Match.png").scaledToHeight(25)
        scan = self.psScan.text()
        self.ndc3 = QLineEdit()
        self.psResult.clear()
        self.psImage.clear()
        sc = str(scan[12:]).lower()
        csc = sc.islower()

        if scan.startswith('M'):
            QMessageBox.warning(self, "Warning", "Not a valid Scan.", QMessageBox.Ok, QMessageBox.Ok)
            self.psScan.clear()
            self.psScan.setFocus()
        else:
            scan = scan.replace('-','')
            if len(scan) > 10 and scan.startswith('003'):
                self.ndc3.setText(scan[3:13])
            elif len(scan) > 10 and scan.startswith('00'):
                self.ndc3.setText(scan[1:])
            elif scan.startswith('3'):
                if len(scan) ==17:
                    self.ndc3.setText(scan[1:12])
                elif len(scan)>10 and len(scan)<17:
                    self.ndc3.setText(scan[1:11])
                elif len(scan)<11:
                    self.ndc3.setText(scan)
            elif scan.startswith('01') and len(scan)>10:
                self.ndc3.setText(scan[5:15])
            elif scan.startswith('17'):
                if sc =='':
                    self.ndc3.setText(scan[0:11])
                elif sc.islower()==True:
                    QMessageBox.warning(self,"Warning", "This is not an NDC barcode.",QMessageBox.Ok,QMessageBox.Ok)
                    self.psScan.clear()
                    self.psScan.setFocus()
            else:
                self.ndc3.setText(scan)
            cps = self.ndc3.text()
            #print(cps)
            param3 = {"labeltype": "all", "query":cps}
            r = requests.get("https://dailymed.nlm.nih.gov/dailymed/search.cfm", params=param3)
            #print(r.url)
            o = urlparse(r.url).query
            q = o.replace("setid=", '')
            #print(o)
            soup2 = bs(r.text, "html5lib")
            #print(soup2)

            dt = soup2.find("td", {"class": "formHeadingTitle"}, string="Packaging")
            # print(dt)

            if dt is None:
                mdt = str(soup2.find("a", {"class": "drug-info-link"}))
                mdt = mdt[mdt.find("setid=") + 6:mdt.find("\">")]
                #print(cps, mdt)
                if mdt == '':
                    restext = " Response from DailyMed: No Drug Package Labels found for PS80 Check."
                    self.psResult.setText(restext)
                    #QMessageBox.warning(self, 'Warning',"No Drug Package Information found.",QMessageBox.Ok, QMessageBox.Ok)
                else:
                    param1 = {"setid": mdt}
                    r1 = requests.get("https://dailymed.nlm.nih.gov/dailymed/drugInfo.cfm", params=param1)
                    # print(r1.url)
                    s2 = bs(r1.text, "html5lib")
                    ps = s2.find_all(string='POLYSORBATE 80')
                    # print(ps)
                    if not ps:
                        ps1 = soup2.find_all(string='Polysorbate 80')
                        # print(ps1,'2')
                        if not ps1:
                            ps2 = soup2.find_all(string='polysorbate 80')
                            if not ps2:
                                restext = ' This product DOES NOT contain Polysorbate 80.'
                                self.psResult.setText(restext)
                                self.psImage.setPixmap(psnf)
                                self.psImage.setMaximumWidth(25)
                            else:
                                QMessageBox.warning(self, 'Information','This product contains ' + str(ps2[0]).capitalize(),QMessageBox.Ok,QMessageBox.Ok)
                                restext = ' This product DOES contain Polysorbate 80.'
                                self.psImage.clear()
                        else:
                            QMessageBox.warning(self, 'Information','This product contains ' + str(ps1[0]).capitalize(),QMessageBox.Ok,QMessageBox.Ok)
                            restext = ' This product DOES contain Polysorbate 80.'
                            self.psImage.clear()

                    else:
                        QMessageBox.warning(self, 'Information', 'This product contains ' + str(ps[0]).capitalize(),QMessageBox.Ok,QMessageBox.Ok)
                        restext = ' This product DOES contain Polysorbate 80.'
                        self.psImage.clear()
            else:
                ps = soup2.find_all(string='POLYSORBATE 80')

                # print(ps,"2")
                if not ps:
                    ps1 = soup2.find_all(string='Polysorbate 80')
                    # print(ps1,'2')
                    if not ps1:
                        ps2 = soup2.find_all(string='polysorbate 80')
                        if not ps2:
                            restext = ' This product DOES NOT contain Polysorbate 80.'
                            self.psResult.setText(restext)
                            self.psImage.setPixmap(psnf)
                            self.psImage.setMaximumWidth(25)
                        else:
                            QMessageBox.warning(self, 'Information','This product contains ' + str(ps2[0]).capitalize(),QMessageBox.Ok,QMessageBox.Ok)
                            restext = ' This product DOES contain Polysorbate 80.'
                            self.psImage.clear()

                    else:
                        QMessageBox.warning(self, 'Information', 'This product contains ' + str(ps1[0]).capitalize(),QMessageBox.Ok,QMessageBox.Ok)
                        restext = ' This product DOES contain Polysorbate 80.'
                        self.psImage.clear()

                else:
                    QMessageBox.warning(self, 'Information', 'This product contains ' + str(ps[0]).capitalize(),QMessageBox.Ok,QMessageBox.Ok)
                    restext = ' This product DOES contain Polysorbate 80.'
                    self.psImage.clear()
            os.chmod(log, S_IWUSR | S_IREAD)
            logw = open(log, "a+")
            logw.write("\nResponse from DailyMed: " + self.ndc3.text() + restext)
            logw.close()
            os.chmod(log, S_IREAD | S_IRGRP | S_IROTH)
            self.ndc3.clear()
            self.psScan.clear()
            self.psScan.setFocus()
        self.timer.start(14400000)

    def submitDisp(self):
        global log
        global sdt
        q2 = "exec KBMADispR @createdwhen = ?, @MMNum = ?, @dose = ?, @NDC = ?, @Lot = ?, @Exp = ?, @User = ?, @ScanStatus = ?;"
        if len(self.lot2.text())==0 or len(self.exp2.text())==0:
            QMessageBox.warning(self, "Warning", "All fields must have a value.", QMessageBox.Ok, QMessageBox.Ok)
            self.scan2.selectAll()
            self.scan2.setFocus()
        elif self.f.count() == 22 and (len(self.lot2.text())==0 or len(self.exp2.text())==0 or len(self.f.itemAt(19).widget().text())==0 or len(self.f.itemAt(21).widget().text())==0):
            QMessageBox.warning(self, "Warning", "All fields must have a value.", QMessageBox.Ok, QMessageBox.Ok)
            self.scan2.selectAll()
            self.scan2.setFocus()
        elif self.f.count() > 22 and (len(self.lot2.text())==0 or len(self.exp2.text())==0 or len(self.f.itemAt(19).widget().text())==0 or len(self.f.itemAt(21).widget().text())==0 or len(self.f.itemAt(25).widget().text())==0 or len(self.f.itemAt(27).widget().text())==0):
            QMessageBox.warning(self, "Warning", "All fields must have a value.", QMessageBox.Ok,QMessageBox.Ok)
            self.scan2.selectAll()
            self.scan2.setFocus()
        elif self.lastScan.text() == "F":
            QMessageBox.warning(self, "Warning", "You can't submit with an Incorrect Scan. Please Scan again.", QMessageBox.Ok, QMessageBox.Ok)
            self.scan2.selectAll()
            self.scan2.setFocus()

        else:
            s1 = [datetime.datetime.now(), self.scan3.text(), self.doseN.text(), self.ndc2.text(), self.lot2.text(), self.exp2.text(), self.user.text(), self.lastScan.text()]
            os.chmod(log, S_IWUSR | S_IREAD)
            logw = open(log, "a+")
            cdisp = pyodbc.connect('Driver={SQL Server};'
                                 'Server=;'
                                # 'Server=;'
                                 'Database=;'
                                # 'Database=;'
                                 'uid=;'
                                 'pwd=;')
                                # 'pwd=;')
            try:
                cd = cdisp.cursor()
                cd.execute(q2, s1)
                cd.commit()

                if self.f.count() == 22:
                    s1[3] = self.f.itemAt(17).widget().text()
                    s1[4] = self.f.itemAt(19).widget().text()
                    s1[5] = self.f.itemAt(21).widget().text()
                    #print(s1)
                    cd.execute(q2, s1)
                    cd.commit()
                elif self.f.count()>22:
                    s1[3] = self.f.itemAt(17).widget().text()
                    s1[4] = self.f.itemAt(19).widget().text()
                    s1[5] = self.f.itemAt(21).widget().text()
                    #print(s1)
                    rows = cd.execute(q2, s1)
                    cd.commit()
                    s1[3] = self.f.itemAt(23).widget().text()
                    s1[4] = self.f.itemAt(25).widget().text()
                    s1[5] = self.f.itemAt(27).widget().text()
                    #print(s1)
                    cd.execute(q2, s1)
                    cd.commit()
            except pyodbc.Error as err:
                logw.write('Error Code 12: SubmitDisp failed! Details: %s\n' % err)
            cd.close()
            logw.write("Dispensing Information submitted for order "+self.srx2.text()+", dose # "+self.doseN.text()+" on "+str(sdt)+" by "+self.user.text()+"\n")
            self.ndc2.clear()
            self.lot2.clear()
            self.exp2.clear()
            self.image_res2.clear()
            self.scan3.clear()
            mNum =''
            item = 16
            while item in range(self.f.count()):
                #item += 1
                self.it = self.f.itemAt(item).widget()
                self.f.removeWidget(self.it)
                self.it.deleteLater()
                self.ndcLine = 4
                del self.it
                #item += 1

            #print(self.f.count())
            self.res.clear()
            self.scan2.clear()
            self.srx2.clear()
            self.srp2.clear()
            self.scan2.setFocus()
            self.lexpi.clear()
            res = []
            self.scan2.setFocus()
        self.timer.start(14400000)
        logw.close()
        os.chmod(log, S_IREAD | S_IRGRP | S_IROTH)

    def rejFind(self):
        global ws
        global log
        rscan = self.rejScan.text()
        rejcnx = pyodbc.connect('Driver={SQL Server};'
                                'Server=;'
                                # 'Server=;'
                                'Database=;'
                                # 'Database=;'
                                'uid=;'
                                'pwd=;')
        # 'pwd=KBMA4Prod;')
        os.chmod(log, S_IWUSR | S_IREAD)
        logw = open(log, "a+")
        if len(self.rejResult.text()) == 0:
            if rscan.startswith("M"):

                self.rejLotLabel.setVisible(False)
                self.rejLotField.setVisible(False)
                self.rejExpLabel.setVisible(False)
                self.rejExpField.setVisible(False)
                self.rejDesc.clear()
                self.rejDoseQty.clear()
                self.rejDoseUom.clear()
                self.rejResult.clear()
                self.rejPtLoc.clear()
                self.rdoseNum.clear()
                if len(rscan) == 10:
                    rmNum = rscan
                    self.scan4.setText(rmNum)
                    # scan3 = mNum
                    rdoseN = None
                    self.rdoseNum.setText(rdoseN)
                elif len(scan) > 10:
                    rmNum = scan[0:10]
                    rdoseN = scan[scan.find(' ') + 1:len(scan)]
                    self.rdoseNum.setText(rdoseN)
                    self.scan4.setText(rmNum)
                res = []

                try:
                    qd = "SET NOCOUNT ON exec MSKKBMA.KBMADisp @MMNum = ?;"
                    rcursor2 = rejcnx.cursor()
                    rcursor2.execute(qd, rmNum)
                except pyodbc.Error as err:
                    logw.write('Error Code 13: RejFind OrderLookup failed! Details: %s\n' % err)
                row = rcursor2.fetchall()
                # print(row)
                if not row:
                    QMessageBox.warning(self, 'Warning', 'The Order scanned is no longer valid.', QMessageBox.Ok,
                                        QMessageBox.Ok)
                    self.rejScan.clear()
                    self.rejScan.setFocus()
                else:
                    self.rejSub.setEnabled(True)
                    for item in row:
                        # print(item[0],item[1])
                        self.rejType.setText("Order")
                        if item[1] is None:
                            orderName = item[0]
                            self.rejResult.setText(item[0])
                        else:
                            orderName = item[0] + ' ' + item[1]
                            self.rejResult.setText(str(item[0]) + ' ' + str(item[1]))
                        # res.append(str(item[2]))
                        # self.res.setText(', '.join(res))
                        self.rejDoseQty.setText(str(item[3]))
                        self.rejDoseUom.setText(str(item[4]))
                        self.rejDesc.setText(str(item[5]))
                        self.rejPtLoc.setText(str(item[6]))
                    # print(self.res.text())
                    # print(self.rejDoseQty.text())
                    # print(self.rejDoseUom.text())
                    # print(self.rejDesc.text())
                    # print(self.rejPtLoc.text())
                    # print(self.srx2.text())
                    rcursor2.close()
            else:
                self.rejType.setText("Batch")
                self.rejLotLabel.setVisible(True)
                self.rejLotField.setVisible(True)
                self.rejExpLabel.setVisible(True)
                self.rejExpField.setVisible(True)
                try:
                    rcursor2 = rejcnx.cursor()
                    qb = "SET NOCOUNT ON EXEC MSKKBMA.KBMAVerifySRx @barcode = ?, @LocationStr = ?;"
                    pb = (rscan, ws)
                    rcursor2.execute(qb, pb)
                except pyodbc.Error as err:
                    QMessageBox.information(self, "Information", 'No results given by the Query', QMessageBox.Ok)
                    # print(err)
                    logw.write('Error Code 14: RejFind VerifySRx failed! Details: %s\n' % err)
                rc6 = rcursor2.fetchall()
                # print(rc6)
                if not rc6:
                    QMessageBox.warning(self, 'Warning',
                                        'This Scancode is not associated to any Products in SRx.\nPlease use the Receiving Tab to send information to PI group.',
                                        QMessageBox.Ok, QMessageBox.Ok)
                    # self.rejResult.setText("Last Scanned Product not in SRx")

                else:
                    self.rejSub.setEnabled(True)
                    self.rejResult.setText(str(rc6[0][1]) + " by " + str(rc6[0][3]))
                    self.rejDoseQty.setText("")
                    self.rejDoseUom.setText("")
                    self.rejDesc.setText("")
                    self.rejPtLoc.setText("")
                rejcnx.close()

        elif self.rejResult.text() != "":
            question = QMessageBox.question(self, 'Message',
                                            "Are you sure you want to cancel this Rejection Reporting?",
                                            QMessageBox.Yes | QMessageBox.No, QMessageBox.Yes)
            if question == QMessageBox.Yes:
                logw.write('Rejection Information for scan ' + self.scan4.text() + ' was not submitted.\n')
                self.rejScan.clear()
                self.rejResult.clear()
                self.rejDesc.clear()
                self.rejDoseQty.clear()
                self.rejDoseUom.setFocus()
                self.rejPtLoc.clear()
                self.rejLotField.clear()
                self.rejExpField.clear()
                self.rejComField.clear()
                self.rejScan.setFocus()
                self.rejSub.setEnabled(False)
            else:
                self.rejScan.undo()
                self.rejScan.undo()
                self.scan2.setFocus()
        self.timer.start(14400000)
        logw.close()
        os.chmod(log, S_IREAD | S_IRGRP | S_IROTH)

    def rejSubmit(self):
        global u
        rejcnx = pyodbc.connect('Driver={ODBC Driver 17 for SQL Server};'
                                # 'Server=;'
                                # 'Server=;'
                                'Server=;'
                                # 'Database=;'
                                # 'Database=;'
                                'Database=;'
                                # 'uid=;'
                                # 'pwd=;')
                                'Trusted_Connection=;')
        # 'pwd=KBMA4Prod;')
        rq = "Insert into RejectionReport VALUES (?,?,?,?,?,?,?,?,?,?,?,?,?,?)"
        rcomm = self.rejComField.toPlainText()

        if self.rejType.text() == "Batch" and (self.rejLotField.text() == "" or self.rejExpField.text() == ""):
            QMessageBox.warning(self, 'Warning', 'Lot and/or Expiration information is missing.', QMessageBox.Ok, QMessageBox.Ok)
        elif self.reasonDrop.currentText() == "Other" and rcomm == "":
            QMessageBox.warning(self, 'Warning', 'Please enter a Comment for this reason.', QMessageBox.Ok, QMessageBox.Ok)
        else:
            '''print(datetime.datetime.now(), u
                  , self.rejType.text(), self.reasonDrop.currentText(),
                  self.rejScan.text(), self.rejResult.text(), self.rejDesc.text(),
                  rcomm,
                  self.rdoseNum.text(), self.rejDoseQty.text(), self.rejDoseUom.text(), self.rejPtLoc.text(),
                  self.rejLotField.text(), self.rejExpField.text())'''
            rp = (datetime.datetime.now(), u
                  , self.rejType.text(), self.reasonDrop.currentText(),
                  self.rejScan.text(), self.rejResult.text(), self.rejDesc.text(),
                  rcomm,
                  self.rdoseNum.text(), self.rejDoseQty.text(), self.rejDoseUom.text(), self.rejPtLoc.text(),
                  self.rejLotField.text(), self.rejExpField.text())
            recrej = rejcnx.cursor()
            recrej.execute(rq, rp)
            recrej.commit()
            recrej.close()
            if self.reasonDrop.currentText() == "Wrong Drug":
                sender = ''
                receiver = ['molinar1@mskcc.org']
                # receiver = 'molinar1@mskcc.org'
                if self.rejType.text() == "Order":
                    message = \
                        """This is a test from my Python app!
    
                        Submitted by: {}@mskcc.org
                        Product Scan: {} 
                        Rejection Type: {} 
                        Rejection Reason: {}
                        Rejected Order: {}
                        Rejected Order Description: {}
                        Order Dose Number: {}
                        Order Dose Qty: {}
                        Order Dose UoM: {}
                        Patient Current Location: {}
                        Comments: {} """.format(u, self.rejScan.text(), self.rejType.text(), self.reasonDrop.currentText(),
                                                self.rejResult.text(), self.rejDesc.text(), self.rdoseNum.text(), self.rejDoseQty.text(),
                                                self.rejDoseUom.text(), self.rejPtLoc.text(), rcomm)
                elif self.rejType.text() == "Batch":
                    message = \
                        """This is a test from my Python app!
    
                        Submitted by: {}@mskcc.org
                        Product Scan: {}
                        Rejection Type: {}
                        Rejection Reason: {}
                        Rejected Product: {}
                        Batch Lot: {}
                        Batch Exp: {}
                        Comments: {} """.format(u, self.rejScan.text(), self.rejType.text(), self.reasonDrop.currentText(),
                                                self.rejResult.text(), self.rejLotField.text(), self.rejExpField.text(), rcomm)

                try:
                    message = MIMEText(message)
                    message['Subject'] = "New Wrong Drug Product Rejection Report"
                    smtpObj = smtplib.SMTP('mskrelay.mskcc.org', 25)
                    smtpObj.sendmail(sender, receiver, message.as_string())
                    # print("sent")
                except Exception as e:
                    os.chmod(log, S_IWUSR | S_IREAD)
                    logw = open(log, "a+")
                    logw.write('Error Code 15: RejSubmit failed! Details: %s\n' % e)
                    logw.close()
                    os.chmod(log, S_IREAD | S_IRGRP | S_IROTH)
            self.rejScan.clear()
            self.rejResult.clear()
            self.rejDesc.clear()
            self.rejDoseQty.clear()
            self.rejDoseUom.setFocus()
            self.rejPtLoc.clear()
            self.rejLotField.clear()
            self.rejExpField.clear()
            self.rejComField.clear()
            self.rejLotLabel.setVisible(False)
            self.rejLotField.setVisible(False)
            self.rejExpLabel.setVisible(False)
            self.rejExpField.setVisible(False)
            self.rejSub.setEnabled(False)
            self.rejScan.setFocus()
        self.timer.start(14400000)

    def Dispensed(self):
        pandas.set_option('display.max_rows', 150000)
        pandas.set_option('display.max_columns', 10)
        pandas.set_option('display.width', 100000)

        cnx = pyodbc.connect('Driver={SQL Server};'
                             'Server=;'
                             'Database=;'
                             'Trusted_Connection=;')

        cursor = cnx.cursor()
        sd = ""
        ed = ""
        list1 = []
        f = open("DispenseReport.txt", 'r')

        q1 = read_sql_query(f.read(), cnx, params=['0', '0', '12'])  # ReDi-Rx Dispensing Query

        start = 10
        df = DataFrame(q1).fillna('')
        # print(df)
        if df.empty:
            # return False
            QMessageBox.information(self, 'Message', "Nothing to Report", QMessageBox.Ok, QMessageBox.Ok)
        else:
            clist = []
            for col in df.columns:
                clist.append(col)
            # print(clist)
            # print(df)
            # creating the document#
            document = Document()
            style = document.styles['Normal']
            font = style.font
            font.name = 'Helvetica'
            font.size = Pt(9)

            # this controls orientation and resizing of page#
            section = document.sections[-1]
            new_width, new_height = section.page_height, section.page_width
            section.orientation = WD_ORIENTATION.LANDSCAPE
            section.page_width = new_width
            section.page_height = new_height

            # adding text to the document#
            h = document.add_paragraph()
            ht = h.add_run('This is a ReDi-Rx Report test')
            ht.font.size = Pt(20)
            ht.bold = True
            h.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

            t = document.add_table(df.shape[0] + 1, df.shape[1], style='Table Grid')

            for row in t.rows:
                for cell in row.cells:
                    paragraphs = cell.paragraphs
                    for paragraph in paragraphs:
                        for run in paragraph.runs:
                            font = run.font
                            font.size = Pt(9)

            for j in range(df.shape[-1]):
                t.cell(0, j).text = df.columns[j]
                # Making Column headers bold#
                ht = t.cell(0, j).paragraphs[0]
                hr = ht.runs
                hf = hr[0].font
                hf.bold = True

            for i in range(df.shape[0]):
                for j in range(df.shape[-1]):
                    t.cell(i + 1, j).text = str(df.values[i, j])
            # pg = document.add_paragraph(str(df))
            document.save('text.docx')
            os.system('text.docx')
            # report = Reporting.Dispensing(self)
            # print(report)
            # if report == False:
            # QMessageBox.information(self,'Message', "Nothing to Report", QMessageBox.Ok, QMessageBox.Ok)
            # else:
            QMessageBox.information(self, 'Message', "Report Printed", QMessageBox.Ok, QMessageBox.Ok)
        self.timer.start(14400000)

#to start the application#
if __name__ == "__main__":
    import sys
    #pyglet.app.run()
    app = QApplication(sys.argv)
    login = loginWindow()
    app.setWindowIcon(QIcon("icon.png"))
    # Set App Icon in Taskbar #
    myApp = "PharmOps"
    ctypes.windll.shell32.SetCurrentProcessExplicitAppUserModelID(myApp)

    # show app in tray icon area #
    #trayIcon = QSystemTrayIcon(QIcon("icon.png"))
    #trayIcon.show()
    if login.exec_() == QDialog.Accepted:
        window = mainWindow()
        window.show()
        sys.exit(app.exec())
